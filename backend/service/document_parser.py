"""
=============================================
文档解析服务
=============================================
支持 PDF, DOCX, TXT, MD, HTML 格式
"""

import os
import re
from typing import Dict, Any, List
from abc import ABC, abstractmethod


# =============================================
# 1. 定义抽象基类（不含 SUPPORTED_TYPES）
# =============================================

class DocumentParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档，返回纯文本内容"""
        raise NotImplementedError

    @classmethod
    def get_parser(cls, file_type: str):
        """获取对应类型的解析器"""
        parser_class = cls.SUPPORTED_TYPES.get(file_type.lower())
        if parser_class is None:
            raise ValueError(f"不支持的文件类型: {file_type}")
        return parser_class()


# =============================================
# 2. 定义所有解析器子类
# =============================================

class PDFParser(DocumentParser):
    """PDF解析器"""

    def parse(self, file_path: str) -> str:
        """使用 pdfplumber 解析 PDF"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return "\n\n".join(text_parts)

        except ImportError:
            try:
                import PyPDF2

                text_parts = []
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

                return "\n\n".join(text_parts)
            except Exception as e:
                raise Exception(f"PDF解析失败: {e}")
        except Exception as e:
            raise Exception(f"PDF解析失败: {e}")


class DOCXParser(DocumentParser):
    """Word文档解析器"""

    def parse(self, file_path: str) -> str:
        """使用 python-docx 解析 Word 文档"""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except Exception as e:
            raise Exception(f"DOCX解析失败: {e}")


class TXTParser(DocumentParser):
    """纯文本解析器"""

    def parse(self, file_path: str) -> str:
        """读取纯文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise Exception(f"TXT文件解析失败: 无法识别编码")


class MarkdownParser(DocumentParser):
    """Markdown解析器"""

    def parse(self, file_path: str) -> str:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        content = re.sub(r'#{1,6}\s+', '', content)
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)
        content = re.sub(r'!\[(.*?)\]\(.*?\)', r'\1', content)
        content = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', content)
        content = re.sub(r'>\s+', '', content)
        content = re.sub(r'-\s+', '', content)
        content = re.sub(r'\*\s+', '', content)
        content = re.sub(r'\d+\.\s+', '', content)

        return content.strip()


class HTMLParser(DocumentParser):
    """HTML解析器"""

    def parse(self, file_path: str) -> str:
        """解析HTML文件"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')

            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()

            text = soup.get_text()
            text = re.sub(r'\n\s*\n', '\n\n', text)

            return text.strip()

        except ImportError:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            content = re.sub(r'<script.*?</script>', '', content, flags=re.DOTALL)
            content = re.sub(r'<style.*?</style>', '', content, flags=re.DOTALL)
            content = re.sub(r'<.*?>', '', content)
            content = re.sub(r'\n\s*\n', '\n\n', content)

            return content.strip()
        except Exception as e:
            raise Exception(f"HTML解析失败: {e}")


# =============================================
# 3. 在基类定义之后添加 SUPPORTED_TYPES
# =============================================

DocumentParser.SUPPORTED_TYPES = {
    "pdf": PDFParser,
    "docx": DOCXParser,
    "txt": TXTParser,
    "md": MarkdownParser,
    "html": HTMLParser
}


# =============================================
# 文本清洗工具
# =============================================

class TextCleaner:
    """文本清洗工具"""

    @staticmethod
    def clean(text: str) -> str:
        """清洗文本"""
        if not text:
            return ""

        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s.,!?;:]', '', text)
        text = re.sub(r'http[s]?://\S+', '', text)
        text = re.sub(r'\S+@\S+', '', text)
        text = text.strip()

        return text

    @staticmethod
    def split_into_sentences(text: str) -> list:
        """将文本分割成句子"""
        sentences = re.split(r'[。！？.!?]', text)
        return [s.strip() for s in sentences if s.strip()]
